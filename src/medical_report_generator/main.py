#!/usr/bin/env python
import sys
import warnings
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import random
from pathlib import Path
from datetime import datetime

from medical_report_generator.crew import MedicalReportGenerator

warnings.filterwarnings("ignore", category=SyntaxWarning, module="pysbd")


def create_word_document_from_template(report_text: str, template_path: str = None, filename: str = "radiology_report.docx"):
    """
    Creates a Word document from the structured report text using a template.
    
    Args:
        report_text: The complete structured report text (output from the crew).
        template_path: Path to the template .docx file. If None, creates from scratch.
        filename: The name of the output .docx file.
    """
    
    # Load template or create new document
    if template_path and Path(template_path).exists():
        print(f"Using template: {template_path}")
        document = Document(template_path)
    else:
        print("Creating document from scratch (template not found or not specified)")
        document = Document()
        # Set default style
        style = document.styles["Normal"] 
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    # Parse the report text to extract sections
    parsed_sections = parse_report_sections(report_text)
    
    if template_path and Path(template_path).exists():
        # Method 1: Replace placeholders in template
        replace_template_placeholders(document, parsed_sections)
    else:
        # Method 2: Build document from scratch (your original method)
        build_document_from_scratch(document, parsed_sections)

    try:
        document.save(filename)
        print(f"\nCompte rendu généré avec succès sous le nom '{filename}'")
        return {"is_generated": True, "filename": filename}
    except Exception as e:
        print(f"\nErreur lors de l'enregistrement du document : {e}")
        return {"is_generated": False, "error": str(e)}


def parse_report_sections(report_text: str) -> dict:
    """
    Parse the report text and extract all sections.
    
    Args:
        report_text: The complete structured report text
        
    Returns:
        dict: Dictionary with section names as keys and content as values
    """
    # Clean up markdown code blocks
    cleaned_text = re.sub(
        r"^\s*```(?:json|text|french)?\s*[\r\n]*(.*?)\s*```\s*$",
        r"\1",
        report_text,
        flags=re.DOTALL | re.IGNORECASE,
    ).strip()

    lines = cleaned_text.split("\n")
    
    # Extract title
    title_content = "Compte Rendu Radiologique"
    lines_for_sections = lines
    
    if lines:
        first_line = lines[0].strip()
        if first_line.upper().startswith("TITRE:"):
            title_content = (
                first_line.split(":", 1)[1].strip()
                if ":" in first_line
                else title_content
            )
            lines_for_sections = lines[1:]

    # Parse sections
    section_headers_list = [
        "Indication:",
        "Technique:", 
        "Incidences:",
        "Résultat:",
        "Conclusion:",
    ]
    
    sections = {"TITRE": title_content}
    current_section_header = None
    current_section_content = []

    for line_idx, line_text in enumerate(lines_for_sections):
        stripped_line = line_text.strip()
        found_new_header = False
        
        for header in section_headers_list:
            if stripped_line.upper().startswith(header.upper()):
                if current_section_header:
                    sections[current_section_header.rstrip(":")] = "\n".join(
                        current_section_content
                    ).strip()
                current_section_header = header
                current_section_content = [stripped_line[len(header):].strip()]
                found_new_header = True
                break
                
        if not found_new_header and current_section_header:
            if stripped_line:
                current_section_content.append(stripped_line)

        if line_idx == len(lines_for_sections) - 1 and current_section_header:
            sections[current_section_header.rstrip(":")] = "\n".join(
                current_section_content
            ).strip()

    # Ensure all required sections exist
    for header in section_headers_list:
        section_name = header.rstrip(":")
        if section_name not in sections:
            sections[section_name] = ""

    return sections


def replace_template_placeholders(document, sections):
    """
    Replace placeholders in the template document with actual content.
    
    This method looks for placeholders like {{TITRE}}, {{Indication}}, etc.
    and replaces them with the actual content.
    """
    # Get current date in French format
    current_date = datetime.now().strftime("%A %d %B %Y")
    # Define placeholder mappings - including all possible variations
    placeholder_mappings = {
        "{{DATE}}": current_date,
        "{{ DATE }}": current_date,
        "{{TITRE}}": sections.get("TITRE", "Compte Rendu Radiologique"),
        "{{ TITRE }}": sections.get("TITRE", "Compte Rendu Radiologique"),
        "{{Indication}}": sections.get("Indication", "Néant"),
        "{{ Indication }}": sections.get("Indication", "Néant"),
        "{{Technique}}": sections.get("Technique", "Néant"),
        "{{ Technique }}": sections.get("Technique", "Néant"),
        "{{Incidences}}": sections.get("Incidences", "Néant"),
        "{{ Incidences }}": sections.get("Incidences", "Néant"),
        "{{Résultat}}": sections.get("Résultat", "Néant"),
        "{{ Résultat }}": sections.get("Résultat", "Néant"),
        "{{Resultat}}": sections.get("Résultat", "Néant"),  # without accent
        "{{ Resultat }}": sections.get("Résultat", "Néant"),  # without accent
        "{{Conclusion}}": sections.get("Conclusion", "Néant"),
        "{{ Conclusion }}": sections.get("Conclusion", "Néant"),
        "{{Conclusions}}": sections.get("Conclusion", "Néant"),  # plural form
        "{{ Conclusions }}": sections.get("Conclusion", "Néant"),  # plural form
        "{{USER}}": "Medical Agent Reporter",  # plural form
    }
    
    # Debug: Print available sections
    print(f"\nDEBUG - Available sections: {list(sections.keys())}")
    for key, value in sections.items():
        print(f"  {key}: {value[:100]}..." if len(str(value)) > 100 else f"  {key}: {value}")
    
    # Function to replace text in paragraph while preserving formatting
    def replace_in_paragraph(paragraph, placeholder, replacement):
        if placeholder in paragraph.text:
            print(f"DEBUG - Found placeholder '{placeholder}' in paragraph")
            
            # Method 1: Try to replace in individual runs first
            replaced = False
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement or "Néant")
                    replaced = True
                    print(f"DEBUG - Replaced '{placeholder}' in run")
            
            # Method 2: If not replaced in runs, rebuild paragraph
            if not replaced and placeholder in paragraph.text:
                print(f"DEBUG - Rebuilding paragraph for '{placeholder}'")
                full_text = paragraph.text
                new_text = full_text.replace(placeholder, replacement or "Néant")
                
                # Clear paragraph and add new text
                paragraph.clear()
                paragraph.add_run(new_text)
    
    # Replace placeholders in paragraphs
    for paragraph in document.paragraphs:
        original_text = paragraph.text
        for placeholder, content in placeholder_mappings.items():
            replace_in_paragraph(paragraph, placeholder, content)
        
        # Log if paragraph changed
        if paragraph.text != original_text:
            print(f"DEBUG - Paragraph changed from: '{original_text[:50]}...' to: '{paragraph.text[:50]}...'")
    
    # Replace placeholders in tables (if your template has tables)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, content in placeholder_mappings.items():
                        replace_in_paragraph(paragraph, placeholder, content)

    # Replace placeholders in headers and footers
    for section in document.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for placeholder, content in placeholder_mappings.items():
                    replace_in_paragraph(paragraph, placeholder, content)
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for placeholder, content in placeholder_mappings.items():
                    replace_in_paragraph(paragraph, placeholder, content)


def build_document_from_scratch(document, sections):
    """
    Build the document from scratch (original method).
    """
    # Add title
    title_paragraph = document.add_paragraph(sections.get("TITRE", "Compte Rendu Radiologique"))
    title_paragraph.style = document.styles["Heading 1"]
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()  # Add a blank line after the title

    # Add sections in order
    section_headers_list = ["Indication", "Technique", "Incidences", "Résultat", "Conclusion"]
    
    for header in section_headers_list:
        content = sections.get(header, "").strip()
        paragraph = document.add_paragraph()
        header_run = paragraph.add_run(f"{header}:")
        header_run.bold = True
        paragraph.add_run(" ")

        if content and content != "-" and content.upper() != "NÉANT":
            paragraph.add_run(content)
        else:
            paragraph.add_run("Néant")
        document.add_paragraph()


# def run(medical_input: str = None):
#     """
#     Run the crew to generate a medical report using template.
#     """
#     print("## Équipe de Génération de Compte Rendu Médical")
#     print("-------------------------------")

#     # Using the French input example for endometriosis
#     raw_medical_input = """
#         Patiente de 38 ans, adressée pour bilan de douleurs pelviennes chroniques avec dysménorrhée et dyspareunie profondes. Suspicion d'endométriose.
#         Antécédents : RAS.
#         Dernières règles : J10 du cycle.
#         L'examen IRM pelvien a été réalisé sur une antenne de 1.5 Tesla, avec des séquences axiales T1, axiales et sagittales T2, axiales T2 avec saturation de graisse, et des séquences de diffusion. Injection de Gadolinium non réalisée.
#         Résultats :
#         Utérus antéversé, de taille normale. Myomètre homogène. Cavité utérine sans particularité. Zone jonctionnelle d'épaisseur normale.
#         Ovaires en position normale, de taille et morphologie conservées, sans kyste suspect.
#         Pas d'anomalie significative au niveau des trompes.
#         Présence de quelques petits foyers d'hypersignal T1 spontané et T2 variable au niveau du cul-de-sac de Douglas et sur le ligament utéro-sacré gauche, compatibles avec des implants d'endométriose superficielle. Le plus large mesure environ 4mm.
#         Pas d'atteinte du torus utérin.
#         Pas de signe d'adénomyose.
#         Pas d'endométriome visible.
#         Pas d'atteinte digestive ou urinaire évidente sur cet examen (limites de la technique pour les atteintes digestives superficielles sans préparation spécifique).
#         Conclusion : Quelques petits implants d'endométriose superficielle visibles au niveau du cul-de-sac de Douglas et du ligament utéro-sacré gauche. Pas de signe d'endométriose profonde ou d'endométriome sur cet examen.
#     """

#     # Define the inputs for the first task
#     inputs = {
#         "raw_input": medical_input if medical_input is not None else raw_medical_input
#     }

#     # Instantiate the Crew using the CrewBase class
#     try:
#         crew_generator = MedicalReportGenerator()
#         crew = crew_generator.crew()

#         # Kick off the crew process
#         print("\nDémarrage du processus de l'équipe...")
#         result = str(crew.kickoff(inputs=inputs))
#         print("\nProcessus de l'équipe terminé.")

#         print("\n## Texte du Compte Rendu Généré:")
#         print(result)
#         print("-------------------------------")

#         # Setup paths
#         current_file_path = Path(__file__).resolve()
#         project_root = current_file_path.parent.parent.parent
        
#         # Template path - adjust this to your template location
#         template_path = project_root / "templates" / "report_template.docx"
        
#         # Output path
#         unique_name = datetime.now().strftime("radiology_report_%Y-%m-%d-%H-%M-%S.docx")
#         generated_report_path_absolute = project_root / "generated" / "reports" / unique_name
#         generated_report_path_relative = Path("generated") / "reports" / unique_name

#         # Generate the .docx file from the final report text using template
#         document_generation_status = create_word_document_from_template(
#             result, 
#             template_path=str(template_path),
#             filename=str(generated_report_path_absolute)
#         )

#         # If successful, replace the absolute path with the relative path string in the return value
#         if document_generation_status["is_generated"]:
#             document_generation_status["filename"] = str(generated_report_path_relative)

#         return document_generation_status

#     except Exception as e:
#         print(
#             f"\nUne erreur s'est produite lors de l'exécution de l'équipe ou de la génération du document : {type(e).__name__}: {e}",
#             file=sys.stderr,
#         )
#         sys.exit(1)



def run(input_file: str = None):
    """
    Run the crew to generate a medical report using template.
    
    Args:
        input_file: Path to the input text file containing medical data.
        If None, will look for files in input_data folder.
    """
    print("## Équipe de Génération de Compte Rendu Médical")
    print("-------------------------------")

    # Setup paths
    current_file_path = Path(__file__).resolve()
    project_root = current_file_path.parent.parent.parent
    input_data_folder = project_root / "input_data"
    
    # Ensure input_data folder exists
    input_data_folder.mkdir(exist_ok=True)
    
    # Determine which input file to use
    if input_file:
        # Use specified file
        input_file_path = Path(input_file)
        if not input_file_path.is_absolute():
            input_file_path = input_data_folder / input_file
    else:
        # Auto-select from input_data folder
        input_files = list(input_data_folder.glob("*.txt"))
        if not input_files:
            print(f"Erreur : Aucun fichier d'entrée trouvé dans {input_data_folder}")
            print("Veuillez créer un fichier .txt dans le dossier input_data/")
            sys.exit(1)
        
        # Use the first available file or let user choose
        if len(input_files) == 1:
            input_file_path = input_files[0]
        else:
            print("Fichiers d'entrée disponibles:")
            for i, file in enumerate(input_files, 1):
                print(f"  {i}. {file.name}")
            
            try:
                choice = int(input("Choisissez un fichier (numéro): ")) - 1
                input_file_path = input_files[choice]
            except (ValueError, IndexError):
                print("Choix invalide. Utilisation du premier fichier.")
                input_file_path = input_files[0]
    
    # Read the input file
    try:
        print(f"Lecture du fichier d'entrée : {input_file_path.name}")
        with open(input_file_path, "r", encoding="utf-8") as f:
            raw_medical_input = f.read().strip()
        
        if not raw_medical_input:
            print("Erreur : Le fichier d'entrée est vide.")
            sys.exit(1)
            
        print(f"Données médicales chargées ({len(raw_medical_input)} caractères)")
        
    except FileNotFoundError:
        print(f"Erreur : Fichier non trouvé : {input_file_path}")
        sys.exit(1)
    except Exception as e:
        print(f"Erreur lors de la lecture du fichier : {e}")
        sys.exit(1)

    # Define the inputs for the first task
    inputs = {
        "raw_input": raw_medical_input
    }

    # Instantiate the Crew using the CrewBase class
    try:
        crew_generator = MedicalReportGenerator()
        crew = crew_generator.crew()

        # Kick off the crew process
        print("\nDémarrage du processus de l'équipe...")
        result = str(crew.kickoff(inputs=inputs))
        print("\nProcessus de l'équipe terminé.")

        print("\n## Texte du Compte Rendu Généré:")
        print(result)
        print("-------------------------------")

        # Template path - adjust this to your template location
        template_path = project_root / "templates" / "report_template.docx"
        
        # Output path with input file name reference
        input_name = input_file_path.stem  # filename without extension
        unique_name = datetime.now().strftime(f"report_{input_name}_%Y-%m-%d-%H-%M-%S.docx")
        generated_report_path_absolute = project_root / "generated" / "reports" / unique_name
        generated_report_path_relative = Path("generated") / "reports" / unique_name

        # Ensure output directory exists
        generated_report_path_absolute.parent.mkdir(parents=True, exist_ok=True)

        # Generate the .docx file from the final report text using template
        document_generation_status = create_word_document_from_template(
            result, 
            template_path=str(template_path),
            filename=str(generated_report_path_absolute)
        )

        # If successful, replace the absolute path with the relative path string in the return value
        if document_generation_status["is_generated"]:
            document_generation_status["filename"] = str(generated_report_path_relative)
            document_generation_status["input_file"] = str(input_file_path.name)

        return document_generation_status

    except Exception as e:
        print(
            f"\nUne erreur s'est produite lors de l'exécution de l'équipe ou de la génération du document : {type(e).__name__}: {e}",
            file=sys.stderr,
        )
        sys.exit(1)




# Keep the other functions as placeholders
def train():
    """Train the crew for a given number of iterations."""
    print("L'entraînement de l'équipe n'est pas complètement implémenté ou testé avec cette configuration LLM.")
    print("Adjust the train function based on CrewAI and LLM provider documentation.")


def replay():
    """Replay the crew execution from a specific task."""
    print("La relecture de l'équipe n'est pas complètement implémentée ou testée avec cette configuration.")
    print("Adjust the replay function based on CrewAI documentation.")


def test():
    """Test the crew execution with sample reports from the testing set."""
    print("## Test du Générateur de Compte Rendu Médical")
    print("-------------------------------")

    current_file_path = Path(__file__).resolve()
    project_root = current_file_path.parent.parent.parent
    test_reports_path = project_root / "knowledge" / "reports" / "testing"
    output_test_reports_path = project_root / "generated" / "testing_outputs"
    output_test_reports_path.mkdir(parents=True, exist_ok=True)

    print(f"Chemin du fichier actuel : {current_file_path}")
    print(f"Racine du projet calculée : {project_root}")
    print(f"Recherche des rapports de test à : {test_reports_path}")

    if not test_reports_path.exists():
        print(f"Erreur : Le répertoire des rapports de test n'a pas été trouvé à {test_reports_path}")
        sys.exit(1)

    test_files = list(test_reports_path.glob("*.txt"))
    if not test_files:
        print(f"Erreur : Aucun fichier de test trouvé dans {test_reports_path}")
        sys.exit(1)

    selected_test_file = random.choice(test_files)
    print(f"\nFichier de test sélectionné : {selected_test_file.name}")

    try:
        with open(selected_test_file, "r", encoding="utf-8") as f:
            ground_truth_report_text = f.read()

        prompt_input = ""
        lines = ground_truth_report_text.splitlines()
        capture = False
        
        for line in lines:
            stripped_line = line.strip()
            if stripped_line.upper().startswith("INDICATION:"):
                prompt_input += stripped_line[len("INDICATION:"):].strip() + " "
                capture = True
            elif capture and not any(
                sh.upper() in stripped_line.upper()
                for sh in ["TECHNIQUE:", "INCIDENCES:", "RÉSULTAT:", "CONCLUSION:"]
            ):
                if stripped_line:
                    prompt_input += stripped_line + " "
            elif capture and any(
                sh.upper() in stripped_line.upper()
                for sh in ["TECHNIQUE:", "INCIDENCES:", "RÉSULTAT:", "CONCLUSION:"]
            ):
                break
                
        prompt_input = prompt_input.strip()

        if not prompt_input:
            print(f"Avertissement : Impossible d'extraire l'indication pour {selected_test_file.name}. Utilisation d'une invite générique.")
            prompt_input = "Patiente consultant pour des douleurs pelviennes et suspicion d'endométriose."

        print(f"\nInvite générée pour l'équipe :\n{prompt_input}")
        print("-------------------------------")

        inputs = {"raw_input": prompt_input}
        crew_generator = MedicalReportGenerator()
        crew = crew_generator.crew()

        print("\nDémarrage du processus de l'équipe pour le test...")
        generated_report_text = str(crew.kickoff(inputs=inputs))
        print("\nProcessus de l'équipe de test terminé.")
        print("\n## Texte du Compte Rendu Généré (Test):")
        print(generated_report_text)
        print("-------------------------------")

        # Use template for test reports too
        template_path = current_file_path.parent.parent.parent / "templates" / "report_template.docx"
        generated_report_filename_docx = output_test_reports_path / f"generated_{selected_test_file.stem}.docx"
        generated_report_filename_txt = output_test_reports_path / f"generated_{selected_test_file.stem}.txt"

        create_word_document_from_template(
            generated_report_text, 
            template_path=str(template_path),
            filename=str(generated_report_filename_docx)
        )

        with open(generated_report_filename_txt, "w", encoding="utf-8") as f:
            f.write("--- INVITE UTILISÉE ---\n")
            f.write(prompt_input + "\n\n")
            f.write("--- COMPTE RENDU GÉNÉRÉ ---\n")
            f.write(generated_report_text + "\n\n")
            f.write("--- COMPTE RENDU DE RÉFÉRENCE (GROUND TRUTH) ---\n")
            f.write(ground_truth_report_text)
            
        print(f"Rapport de test (comparaison) sauvegardé en .txt : {generated_report_filename_txt}")
        print("\n## Comparaison (Manuelle pour l'instant):")
        print(f"Rapport original (vérité terrain) : {selected_test_file.name}")
        print(f"Le rapport généré a été sauvegardé ici : {generated_report_filename_docx}")
        print("Veuillez comparer manuellement le contenu et la structure.")
        print("-------------------------------")

    except FileNotFoundError:
        print(f"Erreur : Le fichier de test {selected_test_file} n'a pas été trouvé.")
    except Exception as e:
        print(f"\nUne erreur s'est produite lors de l'exécution du test : {type(e).__name__}: {e}", file=sys.stderr)




if __name__ == "__main__":
    if len(sys.argv) > 1:
        command = sys.argv[1].lower()
        if command == "run":
            # Check if input file is specified
            input_file = sys.argv[2] if len(sys.argv) > 2 else None
            run(input_file)
        elif command == "train":
            train()
        elif command == "replay":
            replay()
        elif command == "test":
            test()
        else:
            print(f"Commande inconnue : {command}")
            print("Commandes disponibles : run [input_file], test, train, replay")
    else:
        print("Aucune commande fournie. Exécution de la commande 'run' par défaut.")
        run()  # Default command